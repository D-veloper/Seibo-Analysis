# import the necessary data science libraries we need to analyse our data
import pandas as pd
import math
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from sklearn.preprocessing import MinMaxScaler
from docx import Document

pd.set_option('display.max_rows', None)  # set the display option to show all rows when we print our data
pd.set_option('display.max_columns', None)  # set the display option to show all columns when we print our data
pd.set_option('display.max_colwidth', None)  # set the display options to show the full width of the columns

# Load the Excel file. This code assumes the file is in same folder as this script.
file_path = 'Attendance, Enrollment and Meals Fed 2024.xlsx'  # specify the file name.
data = pd.read_excel(file_path)  # read the file and save it inside a variable called data

stat_results = Document()


def print_list_with_and(items):
    # Prints the elements of the list, separated by commas, with 'and' before the last item.
    single = True
    if len(items) > 1:
        # Join all but the last item with commas
        items_string = ", ".join(f'"{item}"' for item in items[:-1])
        # Add 'and' before the last item
        items_string += f', and "{items[-1]}"'
        single = False
        return items_string, single
    elif len(items) == 1:
        # If there is only one item, format it without commas or 'and'
        items_string = f'"{items[0]}"'
        return items_string, single


"""
Get some summary statistics that describe the data and explore relationship between meals fed and attendance
and meals fed and donations.
"""


def get_summary_stats(data, district, count, number_of_districts):
    if district:
        stat_results.add_heading(f"Summary Statistics For {district} district", level=1)
        print(f"\nGenerating Summary stats for {district} district: \n")
        data = data[data['District'] == district]
        summary_data[district] = {}

    description = data.describe()
    # Add a table with the results of data.describe()
    table = stat_results.add_table(rows=description.shape[0] + 1, cols=description.shape[1] + 1)
    # Add header row
    table.cell(0, 0).text = 'Statistic'
    for i, column_name in enumerate(description.columns):
        table.cell(0, i + 1).text = column_name

    # Add statistics row and column names
    for i, index in enumerate(description.index):
        table.cell(i + 1, 0).text = str(index)
        for j, column in enumerate(description.columns):
            table.cell(i + 1, j + 1).text = str(description.at[index, column])

    print(description)
    print(f"\nInterpreting Summary stats for {district} district:\n")

    mean_enrollment_boys = data["Boys Enrolled"].mean()
    mean_attendance_boys = data["Boys Attended"].mean()
    mean_enrollment_girls = data["Girls Enrolled"].mean()
    mean_attendance_girls = data["Girls Attended"].mean()
    mean_enrollment_total = data["Total Enrolled"].mean()
    mean_attendance_total = data["Total Attendance"].mean()
    mean_meals_fed = data["Number of meals fed"].mean()

    std_enrollment_total = data["Total Enrolled"].std()
    std_attendance_total = data["Total Attendance"].std()
    std_meals_fed = data["Number of meals fed"].std()

    # ----------------------------------- Attendance Rate  ---------------------------------------
    sub_heading1 = f"Attendance Rate: {district}"
    stat_results.add_heading(sub_heading1, level=2)
    print(f"----------------------------------- Attendance Rate: {district} ------------------------------------")
    v = (mean_attendance_boys / mean_enrollment_boys) * 100
    attendance_rate_boys = math.ceil(v * 100) / 100

    answer = f"The average attendance rate of male students in {district} district is {attendance_rate_boys}%"
    print(answer)
    stat_results.add_paragraph(answer)
    summary_data[district]["boys_att_rate"] = attendance_rate_boys

    v = (mean_attendance_girls / mean_enrollment_girls) * 100
    attendance_rate_girls = math.ceil(v * 100) / 100

    answer = f"The average attendance rate of female students in {district} district is {attendance_rate_girls}%"
    print(answer)
    stat_results.add_paragraph(answer)
    summary_data[district]["girls_att_rate"] = attendance_rate_girls

    v = (mean_attendance_total / mean_enrollment_total) * 100
    attendance_rate_total = math.ceil(v * 100) / 100

    answer = f"The average attendance rate of students in {district} district is {attendance_rate_total}%"
    print(answer)
    stat_results.add_paragraph(answer)
    summary_data[district]["total_att_rate"] = attendance_rate_total

    # ----------------------------------- Enrolment vs Attendance: Mean ---------------------------------------
    sub_heading2 = f"Mean Enrollment vs Mean Attendance Rate : {district}"
    stat_results.add_heading(sub_heading2, level=2)
    print(f"\n----------------------- Mean Enrollment vs Mean Attendance Rate : {district} -------------------------")
    enrollment_attendance_discrepancy_boys = math.ceil(mean_enrollment_boys - mean_attendance_boys)
    enrollment_attendance_discrepancy_girls = math.ceil(mean_enrollment_girls - mean_attendance_girls)
    enrollment_attendance_discrepancy_total = math.ceil(mean_enrollment_total - mean_attendance_total)

    summary_data[district]["enroll_att_total"] = enrollment_attendance_discrepancy_total
    summary_data[district]["enroll_att_boys"] = enrollment_attendance_discrepancy_boys
    summary_data[district]["enroll_att_girls"] = enrollment_attendance_discrepancy_girls

    discrepancy_result = (f"On average, there is a discrepancy of about {enrollment_attendance_discrepancy_boys} male "
                          f"students between enrollment and attendance in the {district} district")
    print(discrepancy_result)
    stat_results.add_paragraph(discrepancy_result)

    discrepancy_result = (f"On average, there is a discrepancy of about {enrollment_attendance_discrepancy_girls} "
                          f"female students between enrollment and attendance in the {district} district")
    print(discrepancy_result)
    stat_results.add_paragraph(discrepancy_result)

    discrepancy_result = (f"On average, there is a discrepancy of about {enrollment_attendance_discrepancy_total} "
                          f"total students between enrollment and attendance in the {district} district")
    print(discrepancy_result)
    stat_results.add_paragraph(discrepancy_result)

    # ----------------------------------- Enrolment vs Attendance: Variance ---------------------------------------
    sub_heading3 = f"School Meal Impact On Variation in Attendance vs Enrollment : {district}"
    stat_results.add_heading(sub_heading3, level=2)
    print(f"\n---------- School Meal Impact On Variation in Attendance vs Enrollment : {district} -------------")
    summary_data[district]["std_enroll_total"] = std_enrollment_total
    summary_data[district]["std_att_total"] = std_attendance_total

    coefficient_variation_enrollment_total = std_enrollment_total / mean_enrollment_total
    coefficient_variation_attendance_total = std_attendance_total / mean_attendance_total

    cv_comparison = coefficient_variation_attendance_total / coefficient_variation_enrollment_total

    summary_data[district]["cv_comparison"] = cv_comparison

    if cv_comparison > 1.10 or cv_comparison < 0.90:
        if std_enrollment_total > std_attendance_total:
            var_res = (f"The standard deviation for enrollment, {std_enrollment_total}, is higher than for attendance, "
                  f"{std_attendance_total}, which means enrollment numbers fluctuate more than attendance. We compare "
                  f"the coefficient of variation for attendance, {coefficient_variation_attendance_total}, to the "
                  f"coefficient of variation for enrollment, {coefficient_variation_enrollment_total}, to determine how"
                  f" significant this variation is. The coefficient of variation comparison result, {cv_comparison}, "
                  f"shows there is a significant difference of {(abs(1 - cv_comparison)) * 100}% between the "
                  f"enrollment and attendance standard deviations. This means While enrollment varies, the "
                  f"number of students attending is relatively stable. In other words, more students that attend keep "
                  f"attending, indicating the meals we provide might be strong incentives for children to keep "
                  f"attending school. It also demonstrates that providing meals is effective in stabilizing attendance, "
                  f"even if other factors cause enrollment to fluctuate.")
            print(var_res)
            stat_results.add_paragraph(var_res)
        if std_enrollment_total < std_attendance_total:
            var_res = (f"The standard deviation for attendance, {std_attendance_total}, is higher than for enrollment, "
                  f"{std_enrollment_total}, which means attendance numbers fluctuate more than enrollment. We compare "
                  f"the coefficient of variation for attendance, {coefficient_variation_attendance_total}, to the "
                  f"coefficient of variation for enrollment, {coefficient_variation_enrollment_total}, to determine how"
                  f" significant this variation is. The coefficient of variation comparison result, {cv_comparison}, "
                  f"shows a significant difference of {(abs(1 - cv_comparison)) * 100}% between the enrollment and "
                  f"attendance standard deviations. This may indicate that despite efforts to stabilize attendance "
                  f"through providing meals, other factors are causing greater fluctuations in the number of students "
                  f"attending. Meals might provide some incentive, but they may not be sufficient to fully "
                  f"counteract the other challenges students face in consistently attending school.")
            print(var_res)
            stat_results.add_paragraph(var_res)

    elif 0.9 <= cv_comparison or cv_comparison <= 1.10:
        if std_enrollment_total > std_attendance_total:
            var_res = (f"The standard deviation for enrollment, {std_enrollment_total}, is higher than for attendance, "
                  f"{std_attendance_total}, which means enrollment numbers fluctuate more than attendance. We compare "
                  f"the coefficient of variation for attendance, {coefficient_variation_attendance_total}, to the "
                  f"coefficient of variation for enrollment, {coefficient_variation_enrollment_total}, to determine how"
                  f" significant this variation is. The coefficient of variation comparison result, {cv_comparison}, "
                  f"shows that the difference between the enrollment and attendance standard deviations is not "
                  f"significant, with only a {(abs(1 - cv_comparison)) * 100}% difference. This means that while "
                  f"providing meals might play a role to stabilize attendance, the stability in attendance figures "
                  f"is not significantly different from the enrolment figures to distinguish meal feeding as a strong "
                  f"contributing factor to consistency in attendance.")
            print(var_res)
            stat_results.add_paragraph(var_res)
        if std_enrollment_total < std_attendance_total:
            var_res = (f"The standard deviation for attendance, {std_attendance_total}, is higher than for enrollment, "
                  f"{std_enrollment_total}, which means attendance numbers fluctuate more than enrollment. We compare "
                  f"the coefficient of variation for attendance, {coefficient_variation_attendance_total}, to the "
                  f"coefficient of variation for enrollment, {coefficient_variation_enrollment_total}, to determine how"
                  f" significant this variation is. The coefficient of variation comparison result, {cv_comparison}, "
                  f"shows that the difference between the enrollment and attendance standard deviations is not "
                  f"significant, with only a {(abs(1 - cv_comparison)) * 100}% difference. This means that while "
                  f"there is some fluctuation in attendance, it is not drastically different from the fluctuation "
                  f"in enrollment. Therefore, the impact of the meals on attendance might be neutral. "
                  f"Other factors likely contributing to the observed fluctuations in attendance, but these factors "
                  f"are not creating a large enough impact to be considered significant when compared to enrollment "
                  f"variability.")
            print(var_res)
            stat_results.add_paragraph(var_res)

    # ----------------------------------- Meal Distribution ---------------------------------------
    sub_heading4 = f"Consistency of Meal Distribution : {district}"
    stat_results.add_heading(sub_heading4, level=2)
    print(f"\n----------------------- Consistency of Meal Distribution : {district} -------------------------")
    high_consistency_threshold = 0.1 * mean_meals_fed
    mid_consistency_threshold = 0.2 * mean_meals_fed

    if std_meals_fed < high_consistency_threshold:
        meal_const = (f"The standard deviation for the number of meals fed each month is considerably low ({std_meals_fed}) "
              f"relative to the mean ({mean_meals_fed}). This indicates that meal provision and feeding is consistent "
              f"across the district. The schools are consistently feeding around {mean_meals_fed} meals each month, "
              f"with slight fluctuations, which could mean a relatively stable student population and/or meal programme.")
        print(meal_const)
        stat_results.add_paragraph(meal_const)
    elif high_consistency_threshold <= std_meals_fed or std_meals_fed <= mid_consistency_threshold:
        meal_const = (f"The standard deviation for the number of meals fed each month is fairly low ({std_meals_fed}) "
              f"compared to the mean ({mean_meals_fed}). This indicates that meal provision and feeding has some "
              f"variability but remains mostly consistent across the district. The schools are consistently "
              f"feeding around {mean_meals_fed} meals each month, "
              f"with slight fluctuations, which could mean a fairly stable student population and/or meal programme.")
        print(meal_const)
        stat_results.add_paragraph(meal_const)
    elif mid_consistency_threshold < std_meals_fed:
        meal_const = (f"The standard deviation for the number of meals fed each month is large ({std_meals_fed}) "
              f"compared to the mean ({mean_meals_fed}). This indicates that meal provision and feeding has high "
              f"variability and is inconsistent across the district.")
        print(meal_const)
        stat_results.add_paragraph(meal_const)

    # ------------------ Efficiency of Meal Relative to Attendance -----------------------
    sub_heading5 = f"Meal Efficiency Relative to Attendance : {district}"
    stat_results.add_heading(sub_heading5, level=2)
    print(f"\n----------------------- Meal Efficiency Relative to Attendance : {district} -------------------------")
    meals_per_student = math.ceil(mean_meals_fed / mean_attendance_total)
    summary_data[district]["stud_meals_per_month"] = meals_per_student
    v = (meals_per_student / 20) * 100
    meals_per_day = math.ceil(v) / 100
    answer = (f"On Average, each attending student in {district} district receives {meals_per_student} meals per month or "
          f"{meals_per_day} meals per day.")
    print(answer)
    stat_results.add_paragraph(answer)

    if 'Donations' in data.columns:
        # ----------------------------------- Donation vs Meal Fed ---------------------------------------
        sub_heading6 = f"Relationship between meals fed and donation : {district}"
        stat_results.add_heading(sub_heading6, level=2)
        print(f"\n------------------- Relationship between meals fed and donation : {district} -------------------")
        correlation = data['Donations'].corr(data['Number of meals fed'])
        answer = (f"We calculate the correlation coefficient to determine the relationship between donations and meals fed "
              f"in {district} district. The result, {correlation}, shows a ")
        print(answer)
        stat_results.add_paragraph(answer)

        if correlation == 1.0:
            corr_res = ("perfect positive association. This means that more donations perfectly increase the number of meals "
                  "fed. And less donations perfectly decrease the number of meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif 0.8 <= correlation < 1.0:
            corr_res = (
                "very strong positive association. This suggests that as donations increase, the number of meals "
                "fed increases significantly. And as donations decrease, the number of meals fed decreases "
                "significantly.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif 0.6 <= correlation < 0.8:
            corr_res = (
                "strong positive association. This suggests that increased donations are strongly linked to an increase"
                " in the number of meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif 0.4 <= correlation < 0.6:
            corr_res = (
                "moderate positive association. This suggests that there is a moderate increase in the number of meals"
                " fed as donations increase.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif 0.2 <= correlation < 0.4:
            corr_res = ("weak positive association. This suggests that donations have a weak effect on the number of "
                  "meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif 0.0 <= correlation < 0.2:
            corr_res = (
                "very weak positive or no association. This suggests that donations have little to no positive "
                "effect on the number of meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif -0.2 <= correlation < 0.0:
            corr_res = (
                "very weak negative or no association. This suggests that donations have little to no negative "
                "effect on the number of meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif -0.4 <= correlation < -0.2:
            corr_res = (
                "weak negative association. This suggests that as donations increase, the number of meals fed "
                "decreases slightly, but the relationship is weak.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif -0.6 <= correlation < -0.4:
            corr_res = (
                "moderate negative association. This suggests that as donations increase, the number of meals fed "
                "decreases moderately.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif -0.8 <= correlation < -0.6:
            corr_res = (
                "strong negative association. This suggests that increased donations are strongly linked to a "
                "decrease in the number of meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif -1.0 < correlation <= -0.8:
            corr_res = (
                "very strong negative association. This suggests that as donations increase, the number of meals fed "
                "decreases significantly.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)
        elif correlation == -1.0:
            corr_res = ("perfect negative association. This suggests that donations perfectly decrease the number of "
                  "meals fed.")
            print(corr_res)
            stat_results.add_paragraph(corr_res)

        # ----------------------------------- Donation Efficiency ---------------------------------------
        # change donation currency
        sub_heading7 = f"Efficiency of Donations toward Meal Provision : {district}"
        stat_results.add_heading(sub_heading7, level=2)
        print(f"\n------------------ Efficiency of Donations toward Meal Provision : {district} -----------------")
        mean_donation = data["Donations"].mean()
        donation_per_student = mean_donation / mean_attendance_total
        donation_per_meal = mean_donation / mean_meals_fed
        donation_efficiency = math.ceil(donation_per_student / donation_per_meal)

        if donation_efficiency == meals_per_student:
            eff_res = (f"We receive about £{donation_per_student} per student and £{donation_per_meal} per meal. Based on "
                  f"this, we should be able to provide approx {donation_efficiency} meals per student. Since the "
                  f"average meal per student, {meals_per_student}, perfectly matches our expected donation efficiency, "
                  f"{donation_efficiency}, this shows that our feeding programme is financially efficient and w"
                  f"ell managed, with donations being effectively used to feed students.")
            print(eff_res)
            stat_results.add_paragraph(eff_res)
        elif donation_efficiency < meals_per_student:
            if donation_efficiency >= (meals_per_student - (0.1 * meals_per_student)):
                eff_res = (f"We receive about £{donation_per_student} per student and £{donation_per_meal} per meal. Based "
                      f"on this, we should be able to provide approx {donation_efficiency} meals per student. Since "
                      f"the average meal per student, {meals_per_student}, is slightly greater than our expected "
                      f"donation efficiency, {donation_efficiency}, this shows that our feeding programme is providing "
                      f"more meals per student than the donation efficiency would suggest is sustainable. The program "
                      f"may be overextending itself, potentially leading to financial strain. Without increased "
                      f"donations, the program might be unable to maintain this level of meal distribution in the "
                      f"long term.")
                print(eff_res)
                stat_results.add_paragraph(eff_res)
            if donation_efficiency < (meals_per_student - (0.1 * meals_per_student)):
                eff_res = (f"We receive about £{donation_per_student} per student and £{donation_per_meal} per meal. "
                      f"Based on this, we should be able to provide approx {donation_efficiency} meals per student. "
                      f"Since the average meal per student, {meals_per_student}, is significantly greater than our "
                      f"expected donation efficiency, {donation_efficiency}, this shows that our feeding programme "
                      f"is providing more meals per student than might be sustainable given the donation levels. While "
                      f"the current operation is generous, it might not be sustainable without an increase in "
                      f"donations. There is a risk that the program could face financial difficulties if donations do "
                      f"not increase to match the level of service being provided")
                print(eff_res)
                stat_results.add_paragraph(eff_res)
        elif donation_efficiency > meals_per_student:
            if meals_per_student >= (donation_efficiency - (0.1 * donation_efficiency)):
                eff_res = (f"We receive about £{donation_per_student} per student and £{donation_per_meal} per meal. "
                      f"Based on this, we should be able to provide approx {donation_efficiency} meals per student. "
                      f"Since the average meal per student, {meals_per_student}, is slightly less than our expected "
                      f"donation efficiency, {donation_efficiency}, this shows that our feeding programme is providing "
                      f"less meals per student than the donation efficiency would suggest is possible. This points "
                      f"to possible inefficiencies in the program, such as food wastage, misallocation of resources, "
                      f"or operational issues that prevent full utilization of available funds.")
                print(eff_res)
                stat_results.add_paragraph(eff_res)
            if meals_per_student < (donation_efficiency - (0.1 * donation_efficiency)):
                eff_res = (f"We receive about £{donation_per_student} per student and £{donation_per_meal} per meal. "
                      f"Based on this, we should be able to provide approx {donation_efficiency} meals per student. "
                      f"Since the average meal per student, {meals_per_student}, is significantly less than our "
                      f"expected donation efficiency, {donation_efficiency}, this shows that despite sufficient "
                      f"donations, our feeding program is underperforming in terms of meal distribution. Each student "
                      f"should be receiving at least {donation_efficiency} meals based on the donations, but they are "
                      f"only receiving {meals_per_student}. This points to issues with the feeding programme like "
                      f"mismanagement of funds, food wastage, logistical challenges, or other barriers preventing the "
                      f"program from delivering the expected number of meals.")
                print(eff_res)
                stat_results.add_paragraph(eff_res)
    else:
        eff_res = ("Donations column not found in data. Donations-related calculations will be skipped.")
        print(eff_res)
        stat_results.add_paragraph(eff_res)

    if count == number_of_districts:
        sub_heading8 = "Attendance Comparison Across Districts"
        stat_results.add_heading(sub_heading8, level=2)
        print(f"\n------------------------ Attendance Comparison Across Districts -------------------------")
        all_att_rates_boys = []
        all_att_rate_girls = []
        all_att_rates_total = []

        for district in summary_data:
            all_att_rates_boys.append(summary_data[district]["boys_att_rate"])
            all_att_rate_girls.append(summary_data[district]["girls_att_rate"])
            all_att_rates_total.append(summary_data[district]["total_att_rate"])

        max_boys_att_rate = max(all_att_rates_boys)
        min_boys_att_rate = min(all_att_rates_boys)
        max_girls_att_rate = max(all_att_rate_girls)
        min_girls_att_rate = min(all_att_rate_girls)
        max_total_att_rate = max(all_att_rates_total)
        min_total_att_rate = min(all_att_rates_total)

        max_att_district_boys = []
        min_att_district_boys = []
        max_att_district_girls = []
        min_att_district_girls = []
        max_att_district_total = []
        min_att_district_total = []

        # Iterate over summary_data to find the max and min attendance rates
        for dist, stats in summary_data.items():
            boys_att_rate = stats.get("boys_att_rate", 0)
            girls_att_rate = stats.get("girls_att_rate", 0)
            total_att_rate = stats.get("total_att_rate", 0)

            if boys_att_rate == max_boys_att_rate:
                max_att_district_boys.append(dist)
            if boys_att_rate == min_boys_att_rate:
                min_att_district_boys.append(dist)

            if girls_att_rate == max_girls_att_rate:
                max_att_district_girls.append(dist)
            if girls_att_rate == min_girls_att_rate:
                min_att_district_girls.append(dist)

            if total_att_rate == max_total_att_rate:
                max_att_district_total.append(dist)
            if total_att_rate == min_total_att_rate:
                min_att_district_total.append(dist)

        # Print out the results
        result = print_list_with_and(max_att_district_boys)
        if result[1]:
            answer = f"{result[0]} district has the highest attendance rate for boys at {max_boys_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the highest attendance rate for boys at {max_boys_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)

        result = print_list_with_and(min_att_district_boys)
        if result[1]:
            answer = f"{result[0]} district has the lowest attendance rate for boys at {min_boys_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the lowest attendance rate for boys at {min_boys_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)

        result = print_list_with_and(max_att_district_girls)
        if result[1]:
            answer = f"{result[0]} district has the highest attendance rate for girls at {max_girls_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the highest attendance rate for girls at {max_girls_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)

        result = print_list_with_and(min_att_district_girls)
        if result[1]:
            answer = f"{result[0]} district has the lowest attendance rate for girls at {min_girls_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the lowest attendance rate for girls at {min_girls_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)

        result = print_list_with_and(max_att_district_total)
        if result[1]:
            answer = f"{result[0]} district has the highest attendance rate for students at {max_total_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the highest attendance rate for students at {max_total_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)

        result = print_list_with_and(min_att_district_total)
        if result[1]:
            answer = f"{result[0]} district has the lowest attendance rate for students at {min_total_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)
        else:
            answer = f"{result[0]} districts have the lowest attendance rate for students at {min_total_att_rate}%."
            print(answer)
            stat_results.add_paragraph(answer)


# Ensure consistent style
plt.style.use('ggplot')


def get_visualisations(df, district_name):
    """
    Generate relevant visualizations for the given dataframe (df) for the specified district.

    Parameters:
    - df: DataFrame containing attendance and meal data.
    - district_name: String name of the district for which we are generating visuals.
    """
    # Filter for the district
    district_df = df[df['District'] == district_name]

    # Initialize the scaler
    scaler = MinMaxScaler()

    # Extract the relevant columns for scaling
    relevant_columns = district_df[['Number of meals fed', 'Total Attendance', 'Total Enrolled']]

    # Apply Min-Max scaling
    scaled_data = scaler.fit_transform(relevant_columns)

    # Add the scaled data back to the original district_df
    district_df['Scaled Meals Fed'] = scaled_data[:, 0]
    district_df['Scaled Attendance'] = scaled_data[:, 1]
    district_df['Scaled Enrolled'] = scaled_data[:, 2]

    # Set color palette
    base_color = 'grey'
    highlight_color = 'dodgerblue'
    soft_palette = sns.color_palette('pastel')

    # 1. Distribution of Enrollment and Attendance - Histogram and KDE
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    fig.suptitle(f'Distribution of Meals Fed and Attendance for {district_name}', fontsize=16)

    sns.histplot(district_df['Total Enrolled'], kde=True, ax=axes[0], color=highlight_color, bins=20)
    axes[0].set_title('Enrollment Distribution')
    sns.histplot(district_df['Total Attendance'], kde=True, ax=axes[1], color=highlight_color, bins=20)
    axes[1].set_title('Attendance Distribution')
    # sns.histplot(district_df['Scaled Meals Fed'], kde=True, ax=axes[0], color=highlight_color, bins=20)
    # axes[0].set_title('Meals Fed Distribution')

    plt.show()

    # 2. Boxplot of Enrollment and Attendance to show outliers and distribution
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    fig.suptitle(f'Boxplots of Enrollment and Attendance for {district_name}', fontsize=16)

    sns.boxplot(y=district_df['Total Enrolled'], ax=axes[0], color=highlight_color)
    axes[0].set_title('Enrollment Boxplot')

    sns.boxplot(y=district_df['Total Attendance'], ax=axes[1], color=highlight_color)
    axes[1].set_title('Attendance Boxplot')

    # sns.boxplot(y=district_df['Scaled Meals Fed'], ax=axes[2], color=highlight_color)
    # axes[2].set_title('Meals Fed Boxplot')

    plt.show()

    # 3. Enrollment vs. Attendance Over Time
    if 'Month' in district_df.columns:
        plt.figure(figsize=(12, 6))

        plt.plot(district_df['Month'], district_df['Boys Enrolled'], label='Enrolled', color='grey', linewidth=2)
        plt.plot(district_df['Month'], district_df['Boys Attended'], label='Attended', color='#1f77b4',
                 linewidth=2)
        plt.title(f'Boys Enrollment vs. Attendance Over Time: {district} district')
        plt.ylabel('Number of Students')
        plt.legend()
        plt.tight_layout()
        plt.show()

        plt.figure(figsize=(12, 6))
        plt.plot(district_df['Month'], district_df['Girls Enrolled'], label='Enrolled', color='grey',
                 linewidth=2)
        plt.plot(district_df['Month'], district_df['Girls Attended'], label='Attended', color='#1f77b4',
                 linewidth=2)
        plt.title(f'Girls Enrollment vs. Attendance Over Time: {district} district')
        plt.ylabel('Number of Students')
        plt.legend()
        plt.tight_layout()
        plt.show()

        plt.figure(figsize=(12, 6))
        plt.plot(district_df['Month'], district_df['Total Enrolled'], label='Enrolled', color='grey',
                 linewidth=2)
        plt.plot(district_df['Month'], district_df['Total Attendance'], label='Attended', color='#1f77b4',
                 linewidth=2)
        plt.title(f'Total Enrollment vs. Attendance Over Time: {district} district')
        plt.ylabel('Number of Students')
        plt.legend()
        plt.show()

    # 4. Attendance Rates Over Time
    if 'Month' in district_df.columns:
        plt.figure(figsize=(12, 6))
        plt.plot(district_df['Month'], (district_df['Boys Attended'] / district_df['Boys Enrolled']) * 100,
                 label='Boys', color='#1f77b4')
        plt.plot(district_df['Month'], (district_df['Girls Attended'] / district_df['Girls Enrolled']) * 100,
                 label='Girls', color='#ff7f0e')
        plt.plot(district_df['Month'], (district_df['Total Attendance'] / district_df['Total Enrolled']) * 100,
                 label='Total', color='#2ca02c')
        plt.xlabel('Date')
        plt.ylabel('Attendance Rate (%)')
        plt.title(f'Attendance Rates Over Time: {district} District')
        plt.legend()
        plt.tight_layout()
        plt.show()

    # Meals fed vs attendance for boys and girls
    fig, axes = plt.subplots(1, 2, figsize=(18, 6))

    sns.scatterplot(ax=axes[0], x='Number of meals fed', y='Boys Attended', data=district_df, color='#1f77b4')
    sns.regplot(ax=axes[0], x='Number of meals fed', y='Boys Attended', data=district_df, scatter=False, color='grey')
    axes[0].set_xlabel('Number of Meals Fed')
    axes[0].set_ylabel('Boys Attended')
    axes[0].set_title(f'Meals Fed vs. Boys Attendance: {district} District')

    sns.scatterplot(ax=axes[1], x='Number of meals fed', y='Girls Attended', data=district_df, color='#ff7f0e')
    sns.regplot(ax=axes[1], x='Number of meals fed', y='Girls Attended', data=district_df, scatter=False,
                color='grey')
    axes[1].set_xlabel('Number of Meals Fed')
    axes[1].set_ylabel('Girls Attended')
    axes[1].set_title(f'Meals Fed vs. Girls Attendance: {district} District')

    plt.tight_layout()
    plt.show()

    # 2. Scatter plot of Meals Fed vs Attendance for total student population
    plt.figure(figsize=(10, 6))
    sns.scatterplot(x='Number of meals fed', y='Total Attendance', data=district_df, color=highlight_color)
    sns.regplot(x='Number of meals fed', y='Total Attendance', data=district_df, scatter=False, color='grey')
    plt.title(f'Meals Fed vs Attendance in {district_name}')
    plt.xlabel('Meals Fed')
    plt.ylabel('Attendance')
    plt.grid(True)
    plt.show()

    # 4. Correlation heatmap for all numeric variables
    plt.figure(figsize=(8, 6))
    numeric_columns = district_df.select_dtypes(include=np.number)
    corr = numeric_columns.corr()
    sns.heatmap(corr, annot=True, cmap='coolwarm', vmin=-1, vmax=1)
    plt.title(f'Correlation Heatmap for {district_name}')
    plt.show()

    # 5. Line plot to show trends over time (if a date/time feature is present)
    if 'Month' in district_df.columns:
        # Now we can plot with the 'Date' column
        plt.figure(figsize=(10, 6))
        sns.lineplot(x='Month', y='Scaled Meals Fed', data=district_df, color=highlight_color, label='Meals Fed')
        sns.lineplot(x='Month', y='Scaled Attendance', data=district_df, color='lightcoral', label='Attendance')
        plt.title(f'Meals Fed and Attendance Over Time for {district_name}')
        plt.xlabel('Date')
        plt.ylabel('Count')
        plt.legend()
        plt.grid(True)
        plt.show()

    # 5. Donations vs. Meals Fed
    if 'Donations' in district_df.columns:
        plt.figure(figsize=(10, 6))
        sns.scatterplot(x='Donations', y='Number of meals fed', data=district_df, color='#1f77b4')
        sns.regplot(x='Donations', y='Number of meals fed', data=district_df, scatter=False, color='grey')
        plt.xlabel('Donations (£)')
        plt.ylabel('Number of Meals Fed')
        plt.title(f'Donations vs. Meals Fed: {district} District')
        plt.tight_layout()
        plt.show()
    else:
        print("Donations column not found in data. Donations-related visualizations will be skipped.")


summary_data = {}
districts = data['District'].unique()
number_of_districts = data['District'].nunique()
for count, district in enumerate(districts, 1):
    get_summary_stats(data, district, count, number_of_districts)
    get_visualisations(data, district)

stat_results.save("Stat Results.docx")
